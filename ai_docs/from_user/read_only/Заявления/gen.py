import json
import yaml
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# --- Constants for better readability and maintainability ---
# Document margins
MARGIN_TOP = Inches(0.69)
MARGIN_BOTTOM = Inches(0.5)
MARGIN_LEFT = Inches(0.89)
MARGIN_RIGHT = Inches(0.79)

# Font settings
FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)

# Paragraph alignment options
ALIGN_LEFT = WD_PARAGRAPH_ALIGNMENT.LEFT
ALIGN_CENTER = WD_PARAGRAPH_ALIGNMENT.CENTER
ALIGN_RIGHT = WD_PARAGRAPH_ALIGNMENT.RIGHT
ALIGN_JUSTIFY = WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE # Renamed from "width" for clarity

# --- File Loading Functions ---

def load_data_file(filename_base: str, default_data: dict) -> dict:
    """
    Attempts to load data from YAML (.yaml, .yml) or JSON (.json) files
    in that order. Prioritizes YAML.
    """
    possible_filenames = [
        f"{filename_base}.yaml",
        f"{filename_base}.yml",
        f"{filename_base}.json"
    ]

    for filename in possible_filenames:
        try:
            with open(filename, "r", encoding="utf-8") as file:
                if filename.endswith(('.yaml', '.yml')):
                    print(f"Загрузка данных из YAML файла: '{filename}'")
                    return yaml.safe_load(file)
                elif filename.endswith('.json'):
                    print(f"Загрузка данных из JSON файла: '{filename}'")
                    return json.load(file)
        except FileNotFoundError:
            # If file not found, continue to the next possible filename
            continue
        except (yaml.YAMLError, json.JSONDecodeError) as e:
            # If there's a parsing error, print a warning and try the next format/file
            print(f"Ошибка при чтении файла '{filename}': {e}. Попытка следующего формата/файла.")
            continue
        except Exception as e:
            # Catch any other unexpected errors during file loading
            print(f"Неизвестная ошибка при загрузке '{filename}': {e}. Попытка следующего формата/файла.")
            continue

    # If no file was found or successfully loaded after trying all options
    print(f"Предупреждение: Ни один из файлов ({', '.join(possible_filenames)}) не найден или не может быть прочитан. Используются данные по умолчанию.")
    return default_data

def load_config() -> dict:
    """Loads application configuration from 'config.yaml', 'config.yml', or 'config.json'."""
    return load_data_file("config", {
        "director": "",
        "student_name": "",
        "course": "",
        "group": "",
        "specialty": ""
    })

def load_missed_classes() -> dict:
    """Loads missed classes data from 'missed_classes.yaml', 'missed_classes.yml', or 'missed_classes.json'."""
    return load_data_file("missed_classes", {
        "reason": {"instrumental": "", "genitive": ""},
        "document_proof": "",
        "signature_date": "", # Default for the new field
        "lessons": []
    })

# --- DOCX Utility Functions ---

def add_styled_paragraph(doc: Document, text: str, bold: bool = False, align=ALIGN_LEFT):
    """
    Adds a paragraph to the document with standard styling (Times New Roman, 12pt).
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = FONT_SIZE
    run.font.name = FONT_NAME
    p.paragraph_format.line_spacing = 1.0
    p.alignment = align

def add_blank_paragraphs(doc: Document, count: int = 1):
    """Adds multiple blank paragraphs for spacing."""
    for _ in range(count):
        doc.add_paragraph().paragraph_format.line_spacing = 1.0

def setup_document_sections(doc: Document):
    """Sets up standard page size and margins for the document."""
    section = doc.sections[0]
    section.page_height = Inches(11.69) # A4 height
    section.page_width = Inches(8.27)  # A4 width
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

# --- DOCX Content Generation Functions ---

def add_explanatory_note_page(doc: Document, config: dict, missed_data: dict, lesson: dict):
    """Generates a full 'Приложение № 2 (Объяснительная записка)' page."""
    add_styled_paragraph(doc, "Приложение № 2 к Положению", align=ALIGN_RIGHT)
    add_blank_paragraphs(doc, 2)
    add_styled_paragraph(doc, "Директору института (помощнику директора)", align=ALIGN_RIGHT)
    add_styled_paragraph(doc, config.get("director", ""), align=ALIGN_RIGHT)
    add_styled_paragraph(doc, f"от обучающегося {config.get('course', '')} курса, {config.get('group', '')} группы", align=ALIGN_RIGHT)
    add_styled_paragraph(doc, "по специальности (направлению подготовки)", align=ALIGN_RIGHT)
    add_styled_paragraph(doc, config.get("specialty", ""), align=ALIGN_RIGHT)
    add_styled_paragraph(doc, config.get("student_name", ""), align=ALIGN_RIGHT)
    add_blank_paragraphs(doc, 2)

    reason_instrumental = missed_data.get("reason", {}).get("instrumental", "")
    signature_date = missed_data.get("signature_date", "") # Get signature date

    add_styled_paragraph(doc, "Объяснительная записка", bold=True, align=ALIGN_CENTER)
    add_blank_paragraphs(doc)
    add_styled_paragraph(doc, f"Я пропустил(а) занятие {lesson.get('type', '')} по дисциплине {lesson.get('discipline', '')};")
    add_styled_paragraph(doc, f"дата пропуска занятия: {lesson.get('date', '')};")
    add_styled_paragraph(doc, f"в связи с: {reason_instrumental}.")
    add_blank_paragraphs(doc)
    add_styled_paragraph(doc, f"Подтверждающий документ прилагается: {missed_data.get('document_proof', '')}.")
    add_blank_paragraphs(doc, 4)
    # Use the signature_date here
    add_styled_paragraph(doc, f"Дата {signature_date}{' ' * (80 - len(signature_date))} Подпись обучающегося", align=ALIGN_JUSTIFY)
    doc.add_page_break()

def add_application_page(doc: Document, config: dict, missed_data: dict, lesson: dict):
    """Generates a full 'Приложение № 3 (Заявление)' page."""
    add_styled_paragraph(doc, "Приложение № 3 к Положению", align=ALIGN_RIGHT)
    add_blank_paragraphs(doc, 2)
    add_styled_paragraph(doc, "Директору института (помощнику директора)", align=ALIGN_RIGHT)
    add_styled_paragraph(doc, config.get("director", ""), align=ALIGN_RIGHT)
    add_styled_paragraph(doc, f"от обучающегося {config.get('course', '')} курса, {config.get('group', '')} группы", align=ALIGN_RIGHT)
    add_styled_paragraph(doc, "по специальности (направлению подготовки)", align=ALIGN_RIGHT)
    add_styled_paragraph(doc, config.get("specialty", ""), align=ALIGN_RIGHT)
    add_styled_paragraph(doc, config.get("student_name", ""), align=ALIGN_RIGHT)
    add_blank_paragraphs(doc, 2)

    reason_genitive = missed_data.get("reason", {}).get("genitive", "")
    signature_date = missed_data.get("signature_date", "") # Get signature date

    add_styled_paragraph(doc, "Заявление", bold=True, align=ALIGN_CENTER)
    add_blank_paragraphs(doc)
    add_styled_paragraph(doc,
                         f"Прошу не снижать КНЛ, КНС по дисциплине {lesson.get('discipline', '')}, "
                         f"в связи с пропуском занятия по причине: {reason_genitive}")
    add_styled_paragraph(doc, f"дата пропуска занятия: {lesson.get('date', '')}.")
    add_blank_paragraphs(doc, 4)
    # Use the signature_date here
    add_styled_paragraph(doc, f"Дата {signature_date}{' ' * (80 - len(signature_date))} Подпись обучающегося", align=ALIGN_JUSTIFY)
    doc.add_page_break()

# --- Main Document Creation Function ---

def create_docx_documents():
    """
    Main function to load data and generate DOCX documents for each missed lesson.
    """
    config = load_config()
    missed_data = load_missed_classes()

    lessons = missed_data.get("lessons", [])

    if not lessons:
        print("Нет данных о пропущенных занятиях для генерации документов.")
        return

    doc = Document()
    setup_document_sections(doc)

    for i, lesson in enumerate(lessons):
        add_explanatory_note_page(doc, config, missed_data, lesson)
        add_application_page(doc, config, missed_data, lesson)

    current_timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    student_name_sanitized = config.get('student_name', 'Неизвестный').replace(' ', '_').replace('__', '_')
    output_filename = f"Заявление_Объяснительная_{student_name_sanitized}_{current_timestamp}.docx"

    try:
        doc.save(output_filename)
        print(f"Документ '{output_filename}' успешно создан!")
    except Exception as e:
        print(f"Ошибка при сохранении документа: {e}")

if __name__ == "__main__":
    create_docx_documents()

